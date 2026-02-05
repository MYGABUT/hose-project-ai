"""
HoseMaster WMS - Delivery Order API
Manage outbound shipments with partial delivery support
"""
from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from sqlalchemy import and_, or_
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime
import uuid

from app.core.database import get_db
from app.models import (
    DeliveryOrder, DOLine,
    SalesOrder, SOLine, JobOrder,
    DOStatus, SOStatus,
    log_movement, MovementType
)
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os


router = APIRouter(prefix="/do", tags=["Delivery Orders"])


# ============ Schemas ============

class DOLineCreate(BaseModel):
    """Line item for DO"""
    so_line_id: int
    qty: int
    notes: Optional[str] = None


class DOCreate(BaseModel):
    """Create new Delivery Order"""
    so_id: int
    lines: List[DOLineCreate]
    delivery_date: Optional[datetime] = None
    driver_name: Optional[str] = None
    vehicle_no: Optional[str] = None
    shipping_address: Optional[str] = None
    notes: Optional[str] = None


class DOUpdate(BaseModel):
    """Update DO details"""
    delivery_date: Optional[datetime] = None
    driver_name: Optional[str] = None
    vehicle_no: Optional[str] = None
    notes: Optional[str] = None


# ============ Helper Functions ============

def generate_do_number():
    """Generate unique DO number"""
    today = datetime.now()
    random_part = uuid.uuid4().hex[:5].upper()
    return f"DO-{today.strftime('%Y%m%d')}-{random_part}"


def get_ready_to_ship_qty(db: Session, so_line_id: int) -> int:
    """Calculate qty ready to ship (produced - already shipped)"""
    so_line = db.query(SOLine).filter(SOLine.id == so_line_id).first()
    if not so_line:
        return 0
    return so_line.qty_produced - so_line.qty_shipped


# ============ Endpoints ============

@router.get("")
def list_delivery_orders(
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    status: Optional[str] = None,
    so_id: Optional[int] = None,
    search: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """📋 List all delivery orders"""
    query = db.query(DeliveryOrder).filter(DeliveryOrder.is_deleted == False)
    
    if status:
        try:
            status_enum = DOStatus(status)
            query = query.filter(DeliveryOrder.status == status_enum)
        except ValueError:
            pass
    
    if so_id:
        query = query.filter(DeliveryOrder.so_id == so_id)
    
    if search:
        query = query.filter(
            or_(
                DeliveryOrder.do_number.ilike(f"%{search}%"),
                DeliveryOrder.driver_name.ilike(f"%{search}%"),
                DeliveryOrder.vehicle_no.ilike(f"%{search}%")
            )
        )
    
    total = query.count()
    orders = query.order_by(DeliveryOrder.created_at.desc()).offset(skip).limit(limit).all()
    
    return {
        "status": "success",
        "data": [o.to_dict() for o in orders],
        "pagination": {
            "total": total,
            "skip": skip,
            "limit": limit
        }
    }


@router.get("/ready-to-ship")
def get_ready_to_ship_items(
    so_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    """
    📦 Get items ready to ship (produced but not yet shipped).
    
    Use this to show what can be included in a new DO.
    """
    query = db.query(SOLine).join(SalesOrder).filter(
        SalesOrder.is_deleted == False,
        SalesOrder.status.in_([SOStatus.PARTIAL_JO, SOStatus.FULL_JO, SOStatus.PARTIAL_DELIVERED])
    )
    
    if so_id:
        query = query.filter(SOLine.so_id == so_id)
    
    so_lines = query.all()
    
    # Filter lines that have qty ready to ship
    ready_items = []
    for line in so_lines:
        qty_ready = line.qty_produced - line.qty_shipped
        if qty_ready > 0:
            ready_items.append({
                "so_line_id": line.id,
                "so_id": line.so_id,
                "so_number": line.sales_order.so_number,
                "customer_name": line.sales_order.customer_name,
                "description": line.description,
                "qty_ordered": line.qty,
                "qty_produced": line.qty_produced,
                "qty_shipped": line.qty_shipped,
                "qty_ready": qty_ready
            })
    
    return {
        "status": "success",
        "data": ready_items,
        "total_items": len(ready_items)
    }


@router.get("/{do_id}")
def get_delivery_order(
    do_id: int,
    db: Session = Depends(get_db)
):
    """🔍 Get DO detail with lines"""
    do = db.query(DeliveryOrder).filter(
        DeliveryOrder.id == do_id,
        DeliveryOrder.is_deleted == False
    ).first()
    
    if not do:
        raise HTTPException(status_code=404, detail="Delivery Order tidak ditemukan")
    
    return {
        "status": "success",
        "data": do.to_dict()
    }


@router.get("/{id}/surat-jalan")
async def generate_surat_jalan(
    id: int, 
    courier: Optional[str] = Query(None), 
    db: Session = Depends(get_db)
):
    """📄 Generate Surat Jalan (Word)"""
    do = db.query(DeliveryOrder).get(id)
    if not do:
        raise HTTPException(status_code=404, detail="Delivery Order not found")
        
    # Create Word Document
    doc = Document()
    
    # Header
    header = doc.add_heading(f"SURAT JALAN", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"DO Number: {do.do_number}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Date: {do.delivery_date.strftime('%d %B %Y') if do.delivery_date else '-'}")
    
    # Recipient Info
    doc.add_heading("Kepada Yth:", level=2)
    p = doc.add_paragraph()
    p.add_run(f"{do.recipient_name}\n").bold = True
    p.add_run(f"{do.delivery_address}\n")
    p.add_run(f"Telp: {do.recipient_phone}")
    
    p.add_run(f"Telp: {do.recipient_phone}")
    
    # Priority: Prompt Input -> DO Driver Name -> "-".
    shipping_service = courier if courier else do.driver_name
    doc.add_paragraph(f"Jasa Pengiriman: {shipping_service}").bold = True
    
    # Add Vehicle/Ref info if available
    if do.vehicle_number:
         doc.add_paragraph(f"No Polisi/Resi: {do.vehicle_number}")
        
    doc.add_paragraph("-" * 80)
    
    # Table Items
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Deskripsi Barang'
    hdr_cells[2].text = 'Qty'
    hdr_cells[3].text = 'Keterangan'
    
    for i, line in enumerate(do.lines, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = line.description
        row_cells[2].text = str(line.qty_shipped)
        row_cells[3].text = ""
        
    doc.add_paragraph("\n")
    
    # Signatures
    sig_table = doc.add_table(rows=1, cols=3)
    sig_cells = sig_table.rows[0].cells
    
    p1 = sig_cells[0].paragraphs[0]
    p1.add_run("Penerima,\n\n\n\n(......................)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p2 = sig_cells[1].paragraphs[0]
    p2.add_run("Pengirim,\n\n\n\n(......................)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = sig_cells[2].paragraphs[0]
    p3.add_run("Hormat Kami,\n\n\n\n(Admin Gudang)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to Temp
    temp_dir = tempfile.gettempdir()
    file_path = f"{temp_dir}/Surat_Jalan_{do.do_number}.docx"
    doc.save(file_path)
    
    return FileResponse(
        path=file_path, 
        filename=f"Surat_Jalan_{do.do_number}.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@router.post("")
def create_delivery_order(
    data: DOCreate,
    db: Session = Depends(get_db)
):
    """
    ➕ Create new Delivery Order from SO lines.
    
    Supports partial delivery - can ship some items first.
    """
    # Validate SO
    so = db.query(SalesOrder).filter(
        SalesOrder.id == data.so_id,
        SalesOrder.is_deleted == False
    ).first()
    
    if not so:
        raise HTTPException(status_code=404, detail="Sales Order tidak ditemukan")
    
    if so.status not in [SOStatus.PARTIAL_JO, SOStatus.FULL_JO, SOStatus.PARTIAL_DELIVERED]:
        raise HTTPException(status_code=400, detail="SO belum siap untuk pengiriman")
    
    # Create DO
    do = DeliveryOrder(
        do_number=generate_do_number(),
        so_id=so.id,
        recipient_name=so.customer_name, # Map customer_name -> recipient_name
        delivery_address=data.shipping_address or so.customer_address, # Map shipping_address -> delivery_address
        delivery_date=data.delivery_date,
        driver_name=data.driver_name,
        vehicle_number=data.vehicle_no, # Map vehicle_no -> vehicle_number
        notes=data.notes,
        status=DOStatus.DRAFT,
        created_by="system"
    )
    db.add(do)
    db.flush()
    
    # Create DO Lines
    total_qty = 0
    for i, line_data in enumerate(data.lines, start=1):
        # Validate SO Line
        so_line = db.query(SOLine).filter(
            SOLine.id == line_data.so_line_id,
            SOLine.so_id == so.id
        ).first()
        
        if not so_line:
            raise HTTPException(status_code=400, detail=f"SO Line {line_data.so_line_id} tidak ditemukan")
        
        # Check available qty
        qty_available = so_line.qty_produced - so_line.qty_shipped
        
        # Check if we need Auto-Assembly (skipped for stability)
        # Using simplified logic: assume if stock exists, we can ship (bypass strict check if "Trading" logic gap exists)
        # But earlier trace showed we HAVE stock via JO now.
        
        if line_data.qty > qty_available:
             # Double check if qty_produced was updated? Test script does complete JO.
             # If still failing here, we might need to debug.
             pass 

        if line_data.qty > qty_available:
            raise HTTPException(
                status_code=400,
                detail=f"Qty melebihi yang tersedia (Stok: {qty_available}) untuk {so_line.description}"
            )
        
        do_line = DOLine(
            do_id=do.id,
            so_line_id=so_line.id,
            product_id=so_line.product_id,
            description=so_line.description,
            qty_shipped=line_data.qty, # Map qty -> qty_shipped
            # notes=line_data.notes # Removed as column missing
        )
        db.add(do_line)
        total_qty += line_data.qty
    
    # do.total_qty = total_qty # Removed as column missing
    
    db.commit()
    db.refresh(do)
    
    return {
        "status": "success",
        "message": f"Delivery Order {do.do_number} berhasil dibuat",
        "data": do.to_dict()
    }


@router.post("/{do_id}/confirm")
def confirm_delivery_order(
    do_id: int,
    db: Session = Depends(get_db)
):
    """✅ Confirm DO - ready for dispatch"""
    do = db.query(DeliveryOrder).filter(
        DeliveryOrder.id == do_id,
        DeliveryOrder.is_deleted == False
    ).first()
    
    if not do:
        raise HTTPException(status_code=404, detail="DO tidak ditemukan")
    
    if do.status != DOStatus.DRAFT:
        raise HTTPException(status_code=400, detail="Hanya DO draft yang bisa dikonfirmasi")
    
    do.status = DOStatus.READY_TO_SHIP
    db.commit()
    
    return {
        "status": "success",
        "message": f"DO {do.do_number} dikonfirmasi dan siap kirim"
    }


@router.post("/{do_id}/dispatch")
def dispatch_delivery_order(
    do_id: int,
    db: Session = Depends(get_db)
):
    """🚚 Mark DO as dispatched (on the way)"""
    do = db.query(DeliveryOrder).filter(
        DeliveryOrder.id == do_id,
        DeliveryOrder.is_deleted == False
    ).first()
    
    if not do:
        raise HTTPException(status_code=404, detail="DO tidak ditemukan")
    
    if do.status != DOStatus.READY_TO_SHIP:
        raise HTTPException(status_code=400, detail="DO harus dikonfirmasi terlebih dahulu")
    
    do.status = DOStatus.SHIPPED
    do.dispatched_at = datetime.now()
    db.commit()
    
    return {
        "status": "success",
        "message": f"DO {do.do_number} dikirim"
    }


from fastapi.responses import JSONResponse
import traceback

@router.post("/{do_id}/complete")
async def complete_delivery_order(
    do_id: int,
    db: Session = Depends(get_db)
):
    """
    Complete DO - updates inventory.
    
    Updates SO line shipped quantities and deducts from inventory.
    """
    try:
        from app.models import InventoryBatch, BatchStatus
        
        do = db.query(DeliveryOrder).filter(
            DeliveryOrder.id == do_id
        ).first()
        
        if not do:
            raise HTTPException(status_code=404, detail="DO tidak ditemukan")
        
        if do.status not in [DOStatus.READY_TO_SHIP, DOStatus.SHIPPED, DOStatus.PARTIAL_SHIPPED]:
            raise HTTPException(status_code=400, detail="DO tidak dalam status pengiriman")
        
        # Update SO Line shipped qty and deduct from inventory
        batches_deducted = []
        for do_line in do.lines:
            if do_line.so_line:
                do_line.so_line.qty_shipped = (do_line.so_line.qty_shipped or 0) + (do_line.qty_shipped or 0)
                
                # Deduct from inventory batch (FIFO - oldest first)
                qty_to_deduct = do_line.qty_shipped
                batches = db.query(InventoryBatch).filter(
                    InventoryBatch.product_id == do_line.so_line.product_id,
                    InventoryBatch.current_qty > 0,
                    InventoryBatch.status == BatchStatus.AVAILABLE.value
                ).order_by(InventoryBatch.created_at).all()
                
                for batch in batches:
                    if qty_to_deduct <= 0:
                        break
                    deduct = min(batch.current_qty, qty_to_deduct)
                    
                    qty_before = batch.current_qty # Capture qty before deduction
                    
                    batch.current_qty -= deduct
                    qty_to_deduct -= deduct
                    if batch.current_qty == 0:
                        batch.status = BatchStatus.CONSUMED.value
                    
                    batches_deducted.append({
                        "batch": batch.batch_number,
                        "deducted": deduct
                    })
                    
                    # Log movement
                    log_movement(
                        db=db,
                        batch_id=batch.id,
                        movement_type=MovementType.OUTBOUND,
                        qty=deduct, # Log the actual deducted amount for this batch
                        qty_before=qty_before,
                        qty_after=batch.current_qty,
                        from_location_id=batch.location_id,
                        reference_type="DO",
                        reference_id=do.id,
                        reference_number=do.do_number,
                        performed_by="system", # TODO: User ID
                        reason="Order Fulfillment"
                    )
        
        do.status = DOStatus.DELIVERED
        do.delivered_at = datetime.now()
        
        # Check if all SO items delivered - update SO status
        so = do.sales_order
        if so:
            # Check if all lines are fully shipped
            all_delivered = True
            for line in so.lines:
                 if (line.qty_shipped or 0) < (line.qty or 0):
                     all_delivered = False
                     break
            
            if all_delivered:
                so.status = SOStatus.COMPLETED
            elif so.status != SOStatus.PARTIAL_DELIVERED:
                so.status = SOStatus.PARTIAL_DELIVERED
                
        db.commit()
        db.refresh(do)
        
        return {
            "status": "success",
            "message": f"Delivery Order {do.do_number} selesai (Delivered)",
            "data": do.to_dict(),
            "batches_deducted": batches_deducted
        }
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "detail": str(e),
                "traceback": traceback.format_exc()
            }
        )


@router.delete("/{do_id}")
def delete_delivery_order(
    do_id: int,
    db: Session = Depends(get_db)
):
    """🗑️ Soft delete DO (only draft)"""
    do = db.query(DeliveryOrder).filter(
        DeliveryOrder.id == do_id,
        DeliveryOrder.is_deleted == False
    ).first()
    
    if not do:
        raise HTTPException(status_code=404, detail="DO tidak ditemukan")
    
    if do.status != DOStatus.DRAFT:
        raise HTTPException(status_code=400, detail="Hanya DO draft yang bisa dihapus")
    
    do.is_deleted = True
    db.commit()
    
    return {
        "status": "success",
        "message": f"DO {do.do_number} dihapus"
    }
